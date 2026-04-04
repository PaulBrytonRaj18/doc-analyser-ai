"""
DOCX Processing Service using python-docx.
"""

import io
from typing import Tuple

from app.core.logging import get_logger

logger = get_logger(__name__)

# Lazy import for optional dependencies
_docx = None


def _get_docx():
    global _docx
    if _docx is None:
        try:
            from docx import Document

            _docx = Document
        except ImportError:
            logger.warning("python-docx not installed. DOCX processing disabled.")
            _docx = False
    return _docx


class DOCXProcessor:
    """Extract text and structure from DOCX files."""

    def extract_text(self, file_bytes: bytes) -> Tuple[str, dict]:
        """Extract text from DOCX file."""
        metadata = {"num_pages": 0, "file_type": "docx", "has_tables": False}

        Document = _get_docx()
        if not Document:
            return (
                "DOCX processing not available. Please install python-docx.",
                metadata,
            )

        try:
            doc = Document(io.BytesIO(file_bytes))

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            full_text = "\n\n".join(paragraphs)

            tables = doc.tables
            if tables:
                metadata["has_tables"] = True
                full_text += self._extract_tables_text(tables)

            metadata["num_pages"] = max(1, len(paragraphs) // 30)

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            full_text = ""

        return full_text, metadata

    def _extract_tables_text(self, tables) -> str:
        """Extract text from tables."""
        table_texts = []
        for table in tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                table_texts.append("\n".join(rows))

        return "\n\n[Tables]\n" + "\n\n".join(table_texts) if table_texts else ""


docx_processor = DOCXProcessor()
