"""
DOCX text extraction using python-docx
Preserves document structure: headings, paragraphs, tables
"""

import io
from typing import Tuple
from utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extract text from DOCX bytes.
    Returns (extracted_text, paragraph_count)
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(file_bytes))
        sections_text = []
        para_count = 0

        # Extract paragraphs with style awareness
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Format headings distinctly
            if "Heading" in style_name:
                level = _get_heading_level(style_name)
                prefix = "#" * level + " "
                sections_text.append(f"{prefix}{text}")
            else:
                sections_text.append(text)

            para_count += 1

        # Extract tables
        for i, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in row_cells if c)
                if row_text:
                    table_rows.append(row_text)

            if table_rows:
                sections_text.append(f"\n[Table {i + 1}]\n" + "\n".join(table_rows))

        # Extract headers/footers
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        sections_text.insert(0, f"[Header]: {para.text.strip()}")
                        break

        full_text = "\n\n".join(sections_text)
        logger.info(f"DOCX extracted: {para_count} paragraphs, {len(full_text)} chars")
        return full_text, para_count

    except ImportError:
        logger.error("python-docx not installed")
        raise ValueError("python-docx library not available")
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def _get_heading_level(style_name: str) -> int:
    """Extract heading level from style name like 'Heading 1'"""
    try:
        parts = style_name.split()
        if len(parts) >= 2 and parts[-1].isdigit():
            return min(int(parts[-1]), 6)
    except Exception:
        pass
    return 1
